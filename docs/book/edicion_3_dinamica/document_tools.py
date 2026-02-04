#!/usr/bin/env python3
"""
Document Tools - Utilidades para conversión y manipulación de documentos.

Este script proporciona herramientas para:
- Convertir archivos Markdown a Word (.docx)
- Convertir archivos Markdown a texto plano (.txt)
- Concatenar múltiples archivos en uno solo
- Conversión por lotes de directorios completos

Uso:
    python document_tools.py --to-docx archivo.md
    python document_tools.py --to-txt archivo.md
    python document_tools.py --concat "capitulo_*.md" --output libro.md
    python document_tools.py --batch-docx ./directorio
    python document_tools.py --batch-txt ./directorio

Caracteristicas:
- Soporte para tablas Markdown -> Tablas Word
- Soporte para bloques de código
- Preservación de headers y listas


Autor: Generado con asistencia de Claude
Fecha: 2026-01-27
"""

import argparse
import glob
import os
import re
import sys
from pathlib import Path
from typing import List, Optional, Union

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None



def md_to_txt(input_path: str, output_path: Optional[str] = None) -> str:
    """
    Convierte un archivo Markdown a texto plano.
    
    Elimina la sintaxis Markdown mientras preserva la estructura legible
    del contenido.
    
    Args:
        input_path: Ruta al archivo Markdown de entrada
        output_path: Ruta de salida (opcional, por defecto usa mismo nombre con .txt)
    
    Returns:
        Ruta del archivo de salida creado
    """
    input_path = Path(input_path)
    
    if not input_path.exists():
        raise FileNotFoundError(f"No se encontró el archivo: {input_path}")
    
    if output_path is None:
        output_path = input_path.with_suffix('.txt')
    else:
        output_path = Path(output_path)
    
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Procesar el contenido Markdown
    text = content
    
    # Convertir headers a texto con formato
    text = re.sub(r'^#{1,6}\s+(.+)$', r'\1', text, flags=re.MULTILINE)
    
    # Eliminar énfasis (bold, italic)
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)  # bold italic
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)      # bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)          # italic
    text = re.sub(r'___(.+?)___', r'\1', text)        # bold italic alt
    text = re.sub(r'__(.+?)__', r'\1', text)          # bold alt
    text = re.sub(r'_(.+?)_', r'\1', text)            # italic alt
    
    # Eliminar código inline
    text = re.sub(r'`(.+?)`', r'\1', text)
    
    # Convertir links [texto](url) a solo texto
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    
    # Eliminar imágenes ![alt](url)
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
    
    # Limpiar bloques de código
    text = re.sub(r'```[\w]*\n', '', text)
    text = re.sub(r'```', '', text)
    
    # Convertir listas con viñetas
    text = re.sub(r'^[\*\-\+]\s+', '• ', text, flags=re.MULTILINE)
    
    # Limpiar líneas horizontales
    text = re.sub(r'^[\-\*_]{3,}$', '\n' + '─' * 40 + '\n', text, flags=re.MULTILINE)
    
    # Limpiar espacios múltiples
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text.strip())
    
    print(f"✓ Convertido: {input_path.name} → {output_path.name}")
    return str(output_path)


def md_to_docx(input_path: str, output_path: Optional[str] = None) -> str:
    """
    Convierte un archivo Markdown a documento Word (.docx).
    
    Preserva la estructura del documento incluyendo headers, párrafos,
    listas y formato básico.
    
    Args:
        input_path: Ruta al archivo Markdown de entrada
        output_path: Ruta de salida (opcional, por defecto usa mismo nombre con .docx)
    
    Returns:
        Ruta del archivo de salida creado
    
    Raises:
        ImportError: Si python-docx no está instalado
    """
    if Document is None:
        print("Error: Se requiere python-docx. Instálalo con:")
        print("  pip install python-docx")
        sys.exit(1)
        print("Error: Se requiere python-docx. Instálalo con:")
        print("  pip install python-docx")
        sys.exit(1)
    
    input_path = Path(input_path)
    
    if not input_path.exists():
        raise FileNotFoundError(f"No se encontró el archivo: {input_path}")
    
    if output_path is None:
        output_path = input_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)
    
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Configurar estilos básicos
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    
    lines = content.split('\n')
    in_code_block = False
    code_content = []
    
    in_table_block = False
    table_content = []

    in_math_block = False
    math_content = []
    
    for line in lines:
        # Manejar bloques de código
        if line.startswith('```'):

            if in_code_block:
                # Fin del bloque de código
                if code_content:
                    code_para = doc.add_paragraph()
                    code_para.style = 'No Spacing'
                    run = code_para.add_run('\n'.join(code_content))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    code_content = []
                in_code_block = False
            else:
                # Si estábamos en una tabla, terminarla antes de empezar código
                if in_table_block:
                    _process_table_block(doc, table_content)
                    in_table_block = False
                    table_content = []
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue

        # Manejar ecuaciones $$
        # Caso 1: Bloque de una sola línea $$ ... $$
        if line.strip().startswith('$$') and line.strip().endswith('$$') and len(line.strip()) > 4:
            # Procesar inmediatamente
            if in_table_block:
                 _process_table_block(doc, table_content)
                 in_table_block = False
                 table_content = []
            
            math_text = line.strip()[2:-2].strip()
            _process_math_block(doc, [math_text])
            continue
            
        # Caso 2: Bloque multi-línea
        if line.strip().startswith('$$'):
            if in_math_block:
                # Fin del bloque matemático
                if in_table_block:
                     _process_table_block(doc, table_content)
                     in_table_block = False
                     table_content = []
                
                _process_math_block(doc, math_content)
                in_math_block = False
                math_content = []
            else:
                # Inicio del bloque
                # Limpiar buffers de tablas pendientes si las hubiera
                if in_table_block:
                     _process_table_block(doc, table_content)
                     in_table_block = False
                     table_content = []
                in_math_block = True
            continue
            
        if in_math_block:
            math_content.append(line)
            continue

        # Manejar tablas

        if line.strip().startswith('|'):
            if not in_table_block:
                in_table_block = True
                table_content = []
            table_content.append(line)
            continue
        
        # Si estábamos en una tabla y la línea actual no es de tabla
        if in_table_block:
            _process_table_block(doc, table_content)
            in_table_block = False
            table_content = []
            # Continuar procesando la línea actual (header, lista, etc)

        
        # Headers
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = re.sub(r'^#+\s*', '', line)
            text = _clean_markdown_text(text)
            
            if level == 1:
                # Usar Heading 1 en lugar de Title
                para = doc.add_heading(text, level=1)
            else:
                para = doc.add_heading(text, level=min(level, 9))
            continue
        
        # Líneas vacías
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # Listas (Bullets)
        if re.match(r'^[\*\-\+]\s+', line):
            text = re.sub(r'^[\*\-\+]\s+', '', line)
            text = _clean_markdown_text(text)
            # Usar List Bullet
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
            continue
        
        # Listas numeradas - MANUALMENTE para evitar continuación global
        if re.match(r'^\d+\.\s+', line):
            # NO quitamos el número, lo dejamos como texto estático
            # Limpiamos el resto de markdown
            text = _clean_markdown_text(line)
            
            # Crear párrafo con sangría francesa (hanging indent)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            
            _add_formatted_text(p, text)
            continue
        
        # Líneas horizontales
        if re.match(r'^[\-\*_]{3,}$', line):
            para = doc.add_paragraph('─' * 50)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        
        # Párrafos normales
        text = _clean_markdown_text(line)
        if text:
            p = doc.add_paragraph()
            _add_formatted_text(p, text)
    
    # Procesar cualquier bloque pendiente al final
    if in_table_block and table_content:
        _process_table_block(doc, table_content)

    doc.save(output_path)
    print(f"✓ Convertido: {input_path.name} → {output_path.name}")
    return str(output_path)


def _process_table_block(doc, lines: List[str]):
    """
    Procesa un bloque de líneas de tabla Markdown y agrega una tabla al documento.
    Maneja parsing de celdas y estilos básicos.
    """
    if not lines:
        return

    # Parsear filas
    rows = []
    for line in lines:
        # Eliminar pipes externos y limpiar espacios
        content = line.strip()
        if content.startswith('|'): content = content[1:]
        if content.endswith('|'): content = content[:-1]
        
        # Dividir por celdas
        # Nota: Esto es una implementación simple que no maneja pipes escapados \|
        cells = [c.strip() for c in content.split('|')]
        rows.append(cells)

    if len(rows) < 1:
        return

    # Validar si hay fila separadora (header separator)
    # Patrón típico: |---|---| o | :--- | :---: |
    has_header = False
    if len(rows) >= 2:
        separator = rows[1]
        # Verificar si parece separador (solo contiene -, :, |, espacios)
        is_sep = True
        for cell in separator:
            if not re.match(r'^[\s\-:|]+$', cell):
                is_sep = False
                break
        if is_sep:
            has_header = True

    # Preparar datos finales
    final_rows = [rows[0]]  # Header siempre va
    
    if has_header:
        # Saltar la fila separadora
        final_rows.extend(rows[2:])
    else:
        # Si no hay header explícito, todo es data (pero rows[0] ya está añadido)
        final_rows.extend(rows[1:])

    # Normalizar número de columnas
    max_cols = 0
    for r in final_rows:
        max_cols = max(max_cols, len(r))
    
    if max_cols == 0:
        return

    # Crear tabla en Word
    table = doc.add_table(rows=len(final_rows), cols=max_cols)
    try:
        table.style = 'Table Grid'
    except KeyError:
        # Si el estilo no existe (raro en plantillas default), no asignar o usar uno standard
        pass
    
    # Llenar celdas
    for i, row_data in enumerate(final_rows):
        row = table.rows[i]
        
        # Si es header (i==0) y teníamos separador, podemos poner negrita
        # Opcionalmente configurar repetición de header en páginas
        
        for j, cell_text in enumerate(row_data):
            if j < max_cols:
                cell = row.cells[j]
                # Usar el helper para texto con math dentro de tablas también
                # Limpiamos markdown básico
                clean_text = _clean_markdown_text(cell_text)
                
                # Accedemos al párrafo de la celda (el primero)
                p = cell.paragraphs[0]
                _add_formatted_text(p, clean_text)
                
                # Formato básico para headers
                if i == 0 and has_header:
                    for run in p.runs:
                        run.bold = True
    
    doc.add_paragraph()  # Espacio después de la tabla


def _process_math_block(doc, lines: List[str]):
    """Procesa bloques de ecuaciones y aplica formato matemático."""
    text = " ".join(lines).strip()
    if not text:
        return
        
    text = _clean_latex_content(text)
    
    # Crear párrafo centrado con fuente matemática
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(12)
    run.italic = True


def _clean_latex_content(text: str) -> str:
    """Limpia y convierte comandos LaTeX a Unicode."""
    # Limpieza previa de comandos de texto LaTeX
    # Convertir \text{...} -> ...
    text = re.sub(r'\\(?:text|mathrm|textbf|textit)\s*\{([^}]+)\}', r'\1', text)
    
    # Reemplazos básicos de LaTeX a Unicode
    replacements = {
        r'\Delta': 'Δ', r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ',
        r'\theta': 'θ', r'\lambda': 'λ', r'\pi': 'π', r'\sigma': 'σ',
        r'\omega': 'ω', r'\phi': 'φ', r'\mu': 'μ', r'\epsilon': 'ε',
        r'\infty': '∞', r'\times': '×', r'\cdot': '·', r'\div': '÷',
        r'\pm': '±', r'\approx': '≈', r'\neq': '≠', r'\leq': '≤',
        r'\geq': '≥', r'\rightarrow': '→', r'\leftarrow': '←',
        r'\Rightarrow': '⇒', r'\Leftarrow': '⇐',
        # Espaciado
        r'\,': ' ', r'\;': ' ', r'\quad': '    ', r'\qquad': '        ',
        # Limpieza de sintaxis residual
        r'\{': '', r'\}': '',  # Eliminar corchetes de agrupación escapados
        r'{': '', r'}': '',    # Eliminar corchetes de agrupación simples
        r'_': '', r'\^': '',   # Eliminar indicadores de subíndice/superíndice (simplificación)
    }
    
    for latex, char in replacements.items():
        text = text.replace(latex, char)
        
    # Limpieza final de espacios extra
    text = re.sub(r'\s+', ' ', text)
    return text


def _add_formatted_text(paragraph, text: str):
    """
    Agrega texto a un párrafo manejando matemáticas inline ($...$).
    """
    # Patrón para capturar $...$
    # Se usa ( ) para mantener el separador en el split
    parts = re.split(r'(\$[^\$]+\$)', text)
    
    for part in parts:
        if part.startswith('$') and part.endswith('$') and len(part) > 2:
            # Es bloque matemático inline
            content = part[1:-1] # Quitar $
            content = _clean_latex_content(content)
            
            run = paragraph.add_run(content)
            run.font.name = 'Cambria Math'
            run.italic = True
        else:
            # Texto normal
            paragraph.add_run(part)


def _clean_markdown_text(text: str) -> str:

    """Limpia texto de sintaxis Markdown inline."""
    # Eliminar énfasis
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'___(.+?)___', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    
    # Eliminar código inline
    text = re.sub(r'`(.+?)`', r'\1', text)
    
    # Convertir links
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    
    # Eliminar imágenes
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
    
    return text


def concatenate_files(
    file_patterns: Union[str, List[str]],
    output_path: str,
    separator: str = "\n\n---\n\n",
    sort_files: bool = True,
    include_filename_header: bool = False
) -> str:
    """
    Concatena múltiples archivos en uno solo.
    
    Args:
        file_patterns: Ruta(s) o patrón(es) glob de archivos a concatenar
        output_path: Ruta del archivo de salida
        separator: Texto separador entre archivos
        sort_files: Si se deben ordenar los archivos alfabéticamente
        include_filename_header: Si se incluye un header con el nombre del archivo
    
    Returns:
        Ruta del archivo de salida creado
    """
    # Recopilar todos los archivos
    files = []
    
    if isinstance(file_patterns, str):
        file_patterns = [file_patterns]
    
    for pattern in file_patterns:
        matched = glob.glob(pattern)
        files.extend(matched)
    
    # Eliminar duplicados manteniendo orden
    seen = set()
    unique_files = []
    for f in files:
        if f not in seen:
            seen.add(f)
            unique_files.append(f)
    files = unique_files
    
    if not files:
        raise ValueError(f"No se encontraron archivos con los patrones: {file_patterns}")
    
    if sort_files:
        files.sort()
    
    output_path = Path(output_path)
    
    contents = []
    for filepath in files:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        if include_filename_header:
            filename = Path(filepath).name
            header = f"<!-- Archivo: {filename} -->\n\n"
            content = header + content
        
        contents.append(content)
    
    final_content = separator.join(contents)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(final_content)
    
    print(f"✓ Concatenados {len(files)} archivos → {output_path.name}")
    for f in files:
        print(f"  • {Path(f).name}")
    
    return str(output_path)


def batch_convert(
    directory: str,
    output_format: str,
    output_dir: Optional[str] = None,
    pattern: str = "*.md"
) -> List[str]:
    """
    Convierte todos los archivos Markdown de un directorio.
    
    Args:
        directory: Directorio con archivos fuente
        output_format: Formato de salida ('docx' o 'txt')
        output_dir: Directorio de salida (opcional, por defecto el mismo)
        pattern: Patrón glob para filtrar archivos
    
    Returns:
        Lista de rutas de archivos creados
    """
    directory = Path(directory)
    
    if not directory.exists():
        raise NotADirectoryError(f"No se encontró el directorio: {directory}")
    
    if output_dir:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
    else:
        output_dir = directory
    
    files = list(directory.glob(pattern))
    
    if not files:
        print(f"No se encontraron archivos con patrón '{pattern}' en {directory}")
        return []
    
    converter = md_to_docx if output_format == 'docx' else md_to_txt
    created = []
    
    print(f"\nConvirtiendo {len(files)} archivos a .{output_format}...")
    print("-" * 50)
    
    for filepath in sorted(files):
        output_path = output_dir / filepath.with_suffix(f'.{output_format}').name
        try:
            result = converter(str(filepath), str(output_path))
            created.append(result)
        except Exception as e:
            print(f"✗ Error en {filepath.name}: {e}")
    
    print("-" * 50)
    print(f"Completado: {len(created)}/{len(files)} archivos convertidos")
    
    return created


def main():
    """Punto de entrada CLI."""
    parser = argparse.ArgumentParser(
        description='Herramientas de conversión de documentos Markdown',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ejemplos:
  python document_tools.py --to-txt capitulo_01.md
  python document_tools.py --to-docx capitulo_01.md
  python document_tools.py --to-docx capitulo_01.md --output mi_documento.docx
  python document_tools.py --concat "capitulo_*.md" --output libro_completo.md
  python document_tools.py --batch-docx ./
  python document_tools.py --batch-txt ./ --output-dir ./txt_output
        """
    )
    
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument('--to-txt', metavar='ARCHIVO',
                       help='Convertir archivo Markdown a texto plano')
    group.add_argument('--to-docx', metavar='ARCHIVO',
                       help='Convertir archivo Markdown a Word')
    group.add_argument('--concat', metavar='PATRON',
                       help='Concatenar archivos (soporta glob, ej: "cap_*.md")')
    group.add_argument('--batch-txt', metavar='DIRECTORIO',
                       help='Convertir todos los .md de un directorio a .txt')
    group.add_argument('--batch-docx', metavar='DIRECTORIO',
                       help='Convertir todos los .md de un directorio a .docx')
    
    parser.add_argument('--output', '-o', metavar='ARCHIVO',
                        help='Ruta del archivo de salida')
    parser.add_argument('--output-dir', metavar='DIRECTORIO',
                        help='Directorio de salida para conversiones por lotes')
    parser.add_argument('--no-sort', action='store_true',
                        help='No ordenar archivos al concatenar')
    parser.add_argument('--no-headers', action='store_true',
                        help='No incluir headers con nombres de archivo al concatenar')
    parser.add_argument('--separator', default='\n\n---\n\n',
                        help='Separador entre archivos concatenados')
    
    args = parser.parse_args()
    
    try:
        if args.to_txt:
            md_to_txt(args.to_txt, args.output)
        
        elif args.to_docx:
            md_to_docx(args.to_docx, args.output)
        
        elif args.concat:
            if not args.output:
                print("Error: --concat requiere --output para especificar el archivo de salida")
                sys.exit(1)
            concatenate_files(
                args.concat,
                args.output,
                separator=args.separator,
                sort_files=not args.no_sort,
            )
        
        elif args.batch_txt:
            batch_convert(args.batch_txt, 'txt', args.output_dir)
        
        elif args.batch_docx:
            batch_convert(args.batch_docx, 'docx', args.output_dir)
    
    except FileNotFoundError as e:
        print(f"Error: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error inesperado: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
